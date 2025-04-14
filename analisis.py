import csv
import subprocess
from docx import Document
import smtplib
from email.message import EmailMessage
import os

def leer_resumenes(csv_path):
    resumenes = []
    with open(csv_path, encoding="utf-8") as file:
        reader = csv.DictReader(file)
        for fila in reader:
            resumenes.append(fila["resumen"])
    return resumenes

def generar_analisis(resumenes, modelo="mistral"):
    texto_resumenes = "\n\n".join(resumenes)
    prompt = f"""Con base en los siguientes resúmenes de noticias del sector energético colombiano, redacta un análisis coherente de 2 páginas que identifique temas comunes, oportunidades, riesgos y tendencias relevantes para la toma de decisiones estratégicas de una empresa del sector:

{texto_resumenes}

Análisis:"""

    resultado = subprocess.run(
        ["ollama", "run", modelo],
        input=prompt,
        capture_output=True,
        text=True
    )

    return resultado.stdout.strip()

def guardar_en_word(texto, archivo="analisis_sector_energetico.docx"):
    doc = Document()
    doc.add_heading("Análisis del Sector Energético Colombiano", 0)
    for parrafo in texto.split("\n"):
        if parrafo.strip():
            doc.add_paragraph(parrafo.strip())
    doc.save(archivo)
    return archivo

def enviar_por_correo(archivo, destinatario):
    emisor = "tu_correo@gmail.com"
    contraseña = "tu_contraseña_de_aplicacion"

    msg = EmailMessage()
    msg["Subject"] = "Análisis del Sector Energético"
    msg["From"] = emisor
    msg["To"] = destinatario
    msg.set_content("Adjunto encontrarás el análisis generado a partir de las noticias recientes del sector energético.")

    with open(archivo, "rb") as f:
        msg.add_attachment(f.read(), maintype="application", subtype="vnd.openxmlformats-officedocument.wordprocessingml.document", filename=archivo)

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
        smtp.login(emisor, contraseña)
        smtp.send_message(msg)

if __name__ == "__main__":
    resumenes = leer_resumenes("noticias_portafolio_resumen_Mistral.csv")
    analisis = generar_analisis(resumenes)
    archivo_word = guardar_en_word(analisis)
    enviar_por_correo(archivo_word, "juan20150235@gmail.com")
    print("📨 Análisis enviado correctamente.")
    